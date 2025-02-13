import json
import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, Any

def sanitize_filename(filename: str) -> str:
    """
    Sanitizes a string to be used as a valid filename.
    Removes invalid characters and replaces whitespaces with underscores.

    Args:
        filename: The original string to sanitize.
    Returns:
        A sanitized string safe for use as a filename.
    """
    # Remove invalid characters (anything except alphanumerics, dots, underscores, and hyphens)
    sanitized = re.sub(r'[^\w.\-]', '', filename)
    # Replace any sequence of whitespace with an underscore
    sanitized = re.sub(r'\s+', '_', sanitized)
    return sanitized

def render_list(doc: Document, paragraph, list_data):
    """
    Renders a list of items in a Word document paragraph with bullet points.
    
    Args:
        doc: The Word document object
        paragraph: The Word document paragraph to render the list in
        list_data: List of dictionaries containing label and details for each item
    """
    current_paragraph = paragraph
    for item in list_data:
        # Clear the first paragraph's text
        if current_paragraph.text:
            current_paragraph.clear()
            
        # Add bullet point
        p = current_paragraph._element
        pPr = p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), "0")
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), "3")
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        
        # Add label in bold
        run = current_paragraph.add_run(f"{item['label']}: ")
        run.bold = True
        
        # Add details
        current_paragraph.add_run(item['details'])
        
        # Insert new paragraph after the current one
        new_paragraph = doc.add_paragraph()
        current_paragraph._p.addnext(new_paragraph._p)
        current_paragraph = new_paragraph

def create_text_resume(resume_data: Dict[str, Any]) -> str:
    """
    Creates a plain text version of the resume.
    
    Args:
        resume_data: Dictionary containing the resume data
    Returns:
        String containing the formatted text resume
    """
    txt_skills = "\n".join([f"- {item}" for item in resume_data['key_skills']])
    txt_8base = "\n".join([f"- {item['label']}: {item['details']}" for item in resume_data['eight_base']])
    txt_appify = "\n".join([f"- {item['label']}: {item['details']}" for item in resume_data['appify']])
    txt_everteam = "\n".join([f"- {item['label']}: {item['details']}" for item in resume_data['everteam']])
    txt_intalio = "\n".join([f"- {item['label']}: {item['details']}" for item in resume_data['intalio']])

    return f"""
Esteban Felipe
(904) 516 0350 
esteban.felipe@gmail.com

PROFESSIONAL SUMMARY
{resume_data['professional_summary']}

KEY SKILLS
{txt_skills}

VP of Product Management / Director of Product Management | 8base | Feb 2022 – Sep 2024 | Miami, FL
{txt_8base}

Director of Solution Engineering | Appify | May 2021 – Feb 2022 | Campbell, CA
{txt_appify}

Director of Product Management / VP of Solution Engineering / CTO | Everteam | Sep 2015 - Aug 2021 | Boston, MA
{txt_everteam}

Senior Solutions Consultant / Product Manager | Intalio | Nov 2008 - Sep 2015 | Palo Alto, CA
{txt_intalio}
"""

def render_resume(resume_data: Dict[str, Any], analysis_data: Dict[str, Any], store_path: str = "/Users/estebanf/Library/CloudStorage/GoogleDrive-esteban.felipe@gmail.com/My Drive/JH 2025") -> None:
    """
    Renders a resume in both DOCX and TXT formats based on the provided data.
    
    Args:
        resume_data: Dictionary containing the resume content
        analysis_data: Dictionary containing the job analysis data
        store_path: Path where the resume files should be stored
    """
    template_path = 'resume.docx'
    doc = Document(template_path)

    # Process each paragraph in the template
    for paragraph in doc.paragraphs:
        if '{{summary}}' in paragraph.text:
            paragraph.text = resume_data['professional_summary']
        
        if '{{skills}}' in paragraph.text:
            # write a line for each key skill
            paragraph.text = ""
            for skill in resume_data['key_skills']:
                paragraph.add_run(skill + '\n')
        
        if '{{8base}}' in paragraph.text:
            paragraph.text = ""
            render_list(doc, paragraph, resume_data['eight_base'])
        
        if '{{appify}}' in paragraph.text:
            paragraph.text = ""
            render_list(doc, paragraph, resume_data['appify'])
        
        if '{{everteam}}' in paragraph.text:
            paragraph.text = ""
            render_list(doc, paragraph, resume_data['everteam'])
        
        if '{{intalio}}' in paragraph.text:
            paragraph.text = ""
            render_list(doc, paragraph, resume_data['intalio'])

    # Create text version of resume
    str_resume = create_text_resume(resume_data)

    # Create folder and save files
    folder_name = sanitize_filename(f"{analysis_data['company_name']}-{analysis_data['job_title']}")
    savepath = sanitize_filename(f"Esteban_Felipe-{analysis_data['company_name']}-RESUME.docx")
    savepath_txt = sanitize_filename(f"Esteban_Felipe-{analysis_data['company_name']}-RESUME.txt")

    # Create folder
    os.makedirs(os.path.join(store_path, folder_name), exist_ok=True)

    # Save the files
    doc.save(os.path.join(store_path, folder_name, savepath))
    with open(os.path.join(store_path, folder_name, savepath_txt), 'w') as f:
        f.write(str_resume)

if __name__ == "__main__":
    # For backwards compatibility and testing
    with open('inputs/04_resume.json', 'r') as f:
        resume_data = json.load(f)
    with open('inputs/02_analysis_result.json', 'r') as f:
        analysis_data = json.load(f)
    
    render_resume(resume_data, analysis_data)