import json
import subprocess
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers.json import JsonOutputParser
from models import Analysis, CompanyResearch,TopSkills,CoverLetter
import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


template = PromptTemplate.from_template("""
# About you
You are an expert cover letter writer. You excel at creating cover letters that are pleasant to read and create strong alignment with the job post. You are also a skilled writer, so when you craft a cover letter, you don't simply repeat information from the job post, but you create cohesive narratives that help the hiring team appreciate the candidate's talent.
# Your task

Produce a tailored cover letter aligned with the job post and the resume. This letter will include
- An opening paragraph
- A Key strengths section
- A value alignment paragraph
- A closing paragraph

To craft this content, you will leverage the information in the job post, the analysis done on the job post, the talent's background, and the relevant accomplishments identified from the talent's experience.

Your primary strategy in crafting the cover letter is to focus on those job requirements that have strong alignment with the candidate's background, combine two or three accomplishments, and provide a narrative that focuses on how the candidate can satisfy the job success criteria and help the company take advantage of the opportunities ahead or tackle the challenges it faces
# Context

The talent has over 15 years of experience and a strong background in product management, solution engineering, and technology consulting. His hands-on technical skills and business acumen balance well. He has also successfully led customer engagements as a sales and technical leader.

Here are the talent's career highlights:

- Conceptualized, developed, and launched a successful Generative AI product at 8base
- Integrated five enterprise products and expanded them with machine learning capabilities and process automation to create the Information Governance Suite at Everteam, leading to the company's acquisition by Kyocera
- Led the development and go-to-market execution of Intalio|Create, a CRM-based cloud development PaaS, leading to its acquisition by ServiceMax
- Managed multiple iterations of Intalio|BPMS, a process automation platform, leading to its acquisition by Everteam
- Subject matter expert in IT Architecture, Process Automation (BPM), Software Integrations, Low-Code platforms, and AI/ML capabilities
- Led technical presales efforts to enterprise customers in areas
- Combines hands-on technical skills with effective team leadership and successful customer engagements  
## Experience

Here's a list of the talent's most relevant accomplishments in the talent's experience
 ### 8base - Director of Product Management / VP of Product Management

8base was a series A company that created low-code development tools for technical users. It had a backend-as-a-service platform that allowed users to create a GraphQL backend for their applications easily It was a low-code web app development tool for Javascript developers. Later, the company introduced Archie, a generative AI-powered platform to assist entrepreneurs, agencies, and product teams with the requirements to create digital products and the associated planning.

**Accomplishments**
```json

{eightBase_accomplishments}

```
### Appify - Director of Solution Engineering
Appify was a series A company with a low-code platform focused on field services operations.

**Accomplishments**
```json

{appify_accomplishments}

```
### Everteam - Director of Product Management / VP of Solution Engineering / CTO

Everteam was a French-based company with a significant presence in EMEA, providing solutions for document management, content analytics, and records management. Everteam acquired Intalio in 2015

**Accomplishments**

```json

{everteam_accomplishments}

```
### Intalio - Senior Solution Consultant / Product Manager

Intalio was a VC-backed company offering an open-source Business Process Management Suite called Intalio|BPMS. It also launched a CRM-based cloud development platform called Intalio|Create, which later sold its intellectual property to ServiceMax

**Accomplishments**

```json

{intalio_accomplishments}

```

# Instructions
- Anything you write needs to be factual and consistent with the information provided about the talent (background and accomplishments). You can't make up information.
- For the cover letter, follow these instructions:
	- **Opening Paragraph**: Capture attention enthusiastically and align yourself with the role and company. Write a compelling introduction that states the position you're applying for, highlights your relevant experience, and aligns with the company's goals or mission; but do not mention the company goals or mission. It should be 3 to five sentences long. Make sure to include as the seconde sentence something similar to "After looking closely at the responsibilities and contrasting them with my background, it felt like a natural fit." (See samples for other forms)
	- **Core Competencies Section**: Open with an intro paragraph with one concise sentence about the candidate’s expertise, followed by a single transition sentence ending with a colon—no additional elaboration. Then highlight key core competencies and experiences that are key to achieving the success criteria. Summarize 3-5 key strengths or experiences tailored to the job description. You are encouraged to combine the most relevant accomplishments into a compelling narrative of the talent's core competencies in alignment with the role. 
	- **Value Alignment Paragraph**: Explain how the candidate can significantly impact the company. While thinking, answer yourself: (1) How will this role help the company address the challenges ahead? (2) How will this role help this company capitalize on the opportunities? (3) What experiences and competencies does the candidate have that can reassure the reader this candidate is the best choice? Do not call out explicitly challenges, opportunities, or candidate experiences; instead, provide a narrative focused on the candidate's competencies. Do not add lame phrases like "I am eager to contribute..."
	- **Call-to-Action Paragraph**: Close with exactly these two lines: ‘Thank you for your time and consideration. I look forward to further communication."
## Writing instructions
### **1. Overall Writing Goals**
- **Replicate Human-Like Nuance:**  
	 • Emulate natural variability in language by mixing complex, longer sentences with shorter, punchy ones.  
	 • Aim for a high level of “perplexity” (diverse word choice and syntax) and “burstiness” (fluctuations in sentence length and structure) to avoid a monotonous rhythm.
	- Maintain a professional, authoritative, and data-driven voice throughout the text.  
	 • The tone should be neutral yet confident, with an emphasis on technical clarity and business impact.
### **2. Voice and Tone**
- **Professional & Direct:**
	 •Write with clarity and purpose. Avoid unnecessary embellishments or excessive formality—polished but natural writing feels more authoritative than rigid corporate-speak.
	 • Ensure every sentence contributes clear information, using a direct active voice wherever possible.

## **3. Sentence Structure and Style**
- **Variety in Sentence Length and Complexity:**  Combine short, impactful sentences with longer, more detailed ones. **Occasionally, use an abrupt sentence to drive a point home.** Not everything needs a perfect, flowing transition—natural speech has variation.
- **Active and Precise Language:**  Favor active voice to keep the text engaging and direct (e.g., “I led the research process” rather than “The research process was led by me”).  Make a point to avoid as much as possible the usage of passive voice
- **Balanced Imperfections:**  
    - **Embrace intentional imperfection for authenticity.** The goal is not to create robotic, overly polished writing but to mirror how expert professionals naturally communicate.
    - **Introduce slight variations in formality, tone, and structure** to break monotony while maintaining professionalism. Controlled imperfections may include:
        - Occasional sentence fragments for emphasis (e.g., _"And that’s where the challenge lies."_)
        - Short, punchy statements between longer, complex sentences to mimic human speech patterns.
        - Light conversational elements that show personality while maintaining credibility (e.g., _"Let’s break this down."_ instead of _"The following section will explain this in detail."_)
        - **Strategic use of rhetorical questions** to engage the reader (e.g., _"Why does this matter? Because without it, prioritization becomes guesswork."_)
    - **Avoid over-editing to the point of sterility.** A slight rawness in phrasing can make the writing feel more genuine.
### 4. **Achieving Human-Like Nuance Through Controlled Imperfections**
- **Why It Matters:**  Polished text that lacks natural variation can feel overly corporate or robotic. Real experts write and speak with rhythm, emphasis, and occasional informal elements. Controlled imperfections create a voice that is authoritative yet approachable.
- **How to Implement It:**
    - **Vary formality where appropriate.** If explaining a highly technical concept, lean professional. If making a strategic point, a slightly informal tone can add impact.
    - **Occasionally break grammar rules for effect.** Short fragments, contractions, and rhetorical questions add energy. Example: _"Metrics matter. But not at the expense of user trust."_
# Cover letter samples

## Sample 1

### Sample Job post

```
## About the job
**About Anaconda
******Be at the center of AI
**** With more than 45 million users, Anaconda is the most popular operating system for AI providing access to the foundational open-source Python packages used in modern AI, data science, and machine learning through a seamless platform. We pioneered the use of Python for data science, championed its vibrant community, and continue to steward open-source projects that make tomorrow’s innovations possible. Our enterprise-grade solutions enable corporate, research, and academic institutions around the world to harness the power of open source for competitive advantage, groundbreaking research, and a better world. To learn more visit https://www.anaconda.com.

Here is why people love most about working here: We’re not just a company, we’re part of a movement. Our dedicated employees and user community are democratizing data science and creating and promoting open-source technologies for a better world, and our commercial offerings make it possible for enterprise users to leverage the most innovative output from open source in a secure, governed way.

**Summary

** Anaconda is seeking a talented**** Sr. Technical Product Manager**** to drive the evolution of our Python Package Build Platform. This is an excellent opportunity for you to combine your expertise in platform design, automation, and AI tooling with your passion for data science and analytics

**What You’ll Do

**

- Led the technical product roadmap for Anaconda’s Python Package Build Platform, focusing on automation, scalability, and integration with AI-driven tooling.
- Collaborate with engineering teams to develop innovative automation and AI-powered solutions to advance platform capabilities.
- Interact with customers, partners, and suppliers to understand requirements and drive product alignment with market needs.
- Translate complex technical needs into clear, actionable product specifications and requirements.
- Prioritize product improvements, feature requests, and bug fixes based on user feedback, business goals, and strategic impact.
- Drive cross-functional collaboration with sales, marketing, UX, Anaconda Core, and operations to ensure seamless product delivery.
- Analyze and resolve issues impacting the platform’s success autonomously and effectively.
- Stay informed of industry trends, open-source developments, and competitive offerings to shape long-term product strategy
**What You Need**
- 8+ years of product management experience with a strong technical focus.
- Deep understanding of how platforms work, including architecture, scalability, and integration.
- Experience designing and implementing automation systems and AI tooling to streamline processes.
- Demonstrated expertise in Python and open-source ecosystems, particularly in package building and management systems.
- Familiarity with CI/CD pipelines and software lifecycle management.
- Proven ability to operate effectively in a matrixed organization, coordinating across diverse teams.
- Team attitude: “I am not done until WE are done”
- Embody our core values:
- Great People
- Great Product
- Great Performance
- Care deeply about fostering an environment where people of all backgrounds and experiences can flourish

**What Will Make You Stand Out**

- Experience working in a fast-paced startup environment
- Experience in an open-source or data science-focused company.
- Strong background in package environment management or long-term supported product development.
- Expertise in building tools that integrate automation and AI to solve technical challenges.
```
### Cover letter
```
Dear Hiring Manager,

I am eager to submit my application for this Sr. Technical Product Manager at Anaconda, Inc. After carefully reviewing the responsibilities of my professional background and goals, I made the easy decision to apply. I bring over 15 years of experience in product management, solution engineering, and technology consulting; I don't just build products—I ensure they scale and deliver measurable impact. My track record includes launching AI-powered platforms, integrating automation into enterprise products, and driving execution through cross-functional collaboration.

I drive impact in high-stakes technical product roles by blending strategic planning, technical innovation, and cross-functional collaboration. In my resume, as well as on the page below, I've taken the time to highlight some of the specific competencies that have prepared me to be a strong contributor to your team:

-   **Strategic Product Roadmaps aligned with Business Goals:** I've been fortunate enough to steer the creation of detailed roadmaps that break down complex technical challenges into straightforward, actionable plans. By aligning what we want to achieve with market needs, I've helped bring ideas to life—from the initial brainstorming to the launch. My focus has always been on ensuring our efforts meet customer needs and align with our business goals.
-   **AI & Automation: Elevating Product Capabilities:** From the initial idea to bringing it to life, I've worked on incorporating AI solutions that boost performance and smooth operations. Whether improving product features or finding ways to automate tasks, I focus on creating solutions that are both scalable and impactful—helping to give an edge in a competitive market.
-   **Cross-functional leadership: Turning Strategy into Execution:** Ensuring alignment among engineering, design, marketing, and sales teams is essential for driving success. I have built and led diverse teams that focus on transforming our product strategies into tangible outcomes. Open communication and meaningful collaboration are vital, and I actively engage with stakeholders to ensure we are all moving in the same direction.

Anaconda sits at the convergence of open-source governance and enterprise AI, a space that demands innovation and operational excellence. My experience in automation and AI—honed through leading machine learning initiatives and scaling enterprise platforms—positions me to tackle these challenges head-on. I am ready to drive innovation by leveraging automation and AI expertise. Anaconda has a significant opportunity to expand its global presence with innovative, Python-focused solutions for data science teams. My expertise in product execution, AI-driven enhancements, and cross-functional leadership can help Anaconda accelerate plaLet's performance, enhance security in package maAnaconda's, and optimize feature delivery—ultimately driving more substantial adoption and revenue growth.

Thank you for your time and consideration. I look forward to further communication and to continuing the conversation.
```

```

## Sample 2

### Sample Job post

At point.me, we’re on a mission to increase the spending power for millions of people by turning loyalty points into powerful currency. To achieve this vision, we are seeking an experienced Senior Product Manager to lead and elevate our flight search experiences. In this role, you will leverage a deep understanding of customer needs, industry trends, and data-driven insights to design and deliver exceptional search features that drive user satisfaction and revenue growth.

What You'll Do

- Enhance Flight Search Features: Own and continuously improve our flight search experience to align with customer expectations and business goals. Identify and prioritize opportunities to introduce innovative product features and new revenue streams within the search experience.

- Performance Metrics: Define and track key performance metrics to measure success and inform feature requirements.

- Product Strategy and Execution: Drive the end-to-end product development lifecycle, from defining the vision and roadmap to executing prototypes and measuring success at scale.

- Launch 0-1 Features: Develop and scale in-app and email search alerts to engage users effectively and improve retention rates.

- Optimize Features: Develop and execute an A/B testing strategy to optimize features, validate hypotheses, and deliver measurable impact.

- Cross-Functional Collaboration: Work closely with the loyalty team to define business logic, success criteria, and priorities for search features, ensuring alignment with broader engineering and organizational initiatives. Collaborate with engineering, design, sales, and marketing teams to influence new business opportunities and drive the partner agenda.

- Stakeholder Communication: Provide transparency and foster collaboration across teams through excellent communication skills, influencing without authority in complex stakeholder environments.

Who You Are

- Product Management: 5+ years of experience as a Product Manager, with a proven track record of delivering impactful product features.

- Search Product Expertise: 2+ years of experience working on a search product, demonstrating expertise in search technologies, personalization, and user experience.

- Data-Driven: 2+ years of experience with data analysis, capable of deriving actionable insights and driving data-informed decision-making. You have the ability to use data to monitor metrics, define success, and identify gaps.

- Entrepreneurial: Proven experience working with tech teams in high-growth startup environments, demonstrating extreme ownership and prioritizing ruthlessly to achieve organizational goals.

- Problem Solver: Strong ability to break down complex problems into actionable steps, guide products from conception to launch, and measure success throughout the product lifecycle.

- Passionately Curious: You are deeply passionate about delivering exceptional customer experiences through innovative travel search solutions, and are deeply curious about how this functionality fits into the broader product roadmap.

- Collaborative Leader: You bring a collaborative mindset, capable of working cross-functionally and influencing technical and non-technical stakeholders at all levels.

- Strategic: You have a strategic approach to defining vision and success while executing with precision at scale.

This is a unique opportunity to shape the future of our flight search experiences, drive business growth, and make a meaningful impact on our users. If this sounds like you, we’d love to hear from you!

**About Point.me

** At point.me, we are committed to simplifying the loyalty points experience and increasing our customers' spending power by helping them see their points as currency. Join us and be part of a fast-growing company where your work will make a real impact. We offer competitive salaries, meaningful equity, comprehensive health coverage, and the opportunity to work fully remotely as part of a distributed team.

Benefits & perks

Join Our Growing Team! Here At Point.me We Believe In Taking Care Of Our Team, So That Our Team Can Take Care Of Our Customers. All Employees Are Offered The Following

- Competitive salaries and meaningful equity

- Comprehensive health care coverage

- A 100% distributed workforce, so you can contribute from wherever you prefer

- An open vacation policy, with a minimum of 15 days off each year

- Team trips and outings

This is a management role, with commensurate responsibility and compensation. The anticipated salary range is $150,000 to $210,000. Specific compensation will be determined based on your level of experience.

Given the nature of our business, we anticipate that all team members will be traveling from time to time, including company trips and off-sites, or meeting with strategic partners. As such, being fully vaccinated against COVID-19 is a condition of employment at point.me. We’ll ask you to send us a copy of your vaccination card in advance of your first day.

While our team members can participate from anywhere, all applicants must be eligible to work in the United States.

We respect the individual needs of employees and are an equal opportunity employer.

```

### Cover letter

```

Dear Hiring Manager,

After carefully reviewing the responsibilities of the Senior Product Manager, Search position at point.me and aligning them with my professional background and goals, I made the easy decision to apply. I bring over 15 years of experience in product management, solution engineering, and technology consulting, with a track record of leveraging data-driven insights to build scalable, customer-centric solutions.

My expertise lies at the intersection of deep technical fluency and strategic execution, enabling me to drive innovation in search and related technologies. In my resume, as well as on the page below, I've taken the time to highlight some of the specific competencies that have prepared me to be a strong contributor to your team:

-   **Data-Driven Product Development:** I've consistently leveraged performance metrics and customer analytics to shape product roadmaps. At 8base, I aligned cross-functional teams around tools like FullStory and Amplitude, accelerating iteration cycles and boosting adoption rates through data-backed decisions.
-   **AI-Powered Search & Optimization:** At Everteam, I led the integration of Apache Solr-powered indexing, NLP-based metadata extraction, and ML-driven content classification to enhance enterprise search capabilities. These innovations improved efficiency and engagement—expertise I can apply to elevate point. point.me's flight search experience.
-   **0-to-1 Feature Execution:** I have successfully led AI-powered product launches, balancing rapid prototyping with user-centric prioritization. My approach ensures early momentum, validated learning, and sustainable growth.
-   **Cross-Functional Leadership:** I've unified globally distributed teams across product, engineering, and design to deliver complex solutions on time. I drive execution by fostering open communication and aligning stakeholder priorities with technical feasibility and strategic goals.

By applying structured experimentation and focusing on personalization, I can drive point.me to optimize search experiences. My experience managing stakeholder relationships also enables me to navigate complex partnerships and align teams around a shared roadmap vision.

Thank you for your time and consideration. I look forward to further communication.
```

## Sample 3

### Sample Job post

```

** We are passionate about delivering high-quality, actionable data to solve real-world challenges in the home services industry. As businesses increasingly rely

We are driven to solve real-world challenges in the home services industry with innovative, AI-powered solutions. As the industry evolves, we seek AI-focused Product Managers who thrive on change and build adaptable, intelligent products. Leveraging TitanIntelligence (TI), our branded AI, we transform data into actionable insights, automation, and decision intelligence that enhance efficiency and performance. Working closely with Product Managers and business stakeholders, we uncover new opportunities to modernize operations, optimize workflows, and help customers lead in their markets.

**Principal TI Data Product Manager

**The Principal TI Data Product Manager is a strategic and hands-on role, responsible for leading AI innovation by identifying opportunities and collaborating with Product Managers and Engineering to prototype and launch AI-driven solutions for real-world challenges. This role requires a strong ability to identify AI opportunities within business processes and to develop complex models into intelligent, scalable solutions with significant impact. Additionally, this role bridges traditional product management, AI engineering, and business strategy to deliver cutting-edge AI solutions that leverage the latest advancements to transform decision-making, automation, and operational efficiency, driving measurable business outcomes.

**As a Principal TI Data Product Manager, you will:

**

- Spearhead the design and implementation of AI-driven data strategies that align with current business objectives and anticipate future needs.

- Identify and translate AI and machine learning models into scalable enterprise solutions.

- Collaborate with cross-functional teams, including Engineering, Customer Success, Product, Legal, and UX, to develop AI solutions that enhance user experience and operational efficiency.

- Design and assess experiments focused on feature development, AI adoption, and data-driven decision-making to optimize outcomes.

- Partner with business stakeholders to offer expertise in AI experimentation, causal inference, and model evaluation, utilizing methods like A/B testing, multivariate testing, and uplift modeling.

- Define, prototype, and validate ML and AI solutions, ensuring feasibility and impact.

- Define model requirements, encompassing functional specifications for LLMs, retrieval augmentation, prompt engineering, and AI-powered automation.

- Work with prospects and customers to identify AI-driven business needs, pinpoint challenges, and showcase capabilities.

- Stay ahead of AI advancements by leveraging techniques such as LLM fine-tuning, retrieval optimization, model customization, and AI performance evaluation.

****To be successful in this role, you'll need:

****

- Passion for building AI-driven products that deliver measurable business impact.

- Experience in AI applications across Fintech, Marketing, or other data-intensive industries.

- A 4-year degree (or equivalent experience) in Computer Science, Engineering, Mathematics, or a related discipline.

- 5+ years of experience in two or more of the following areas:

- AI & Machine Learning Product Development

- AI Experimentation & Model Evaluation

- Technical Product Management or AI Product Management

- Data Science & Analytics

- Strong verbal and written communication skills, with experience influencing cross-functional teams.

- An entrepreneurial mindset with the ability to work independently and drive AI initiatives.

- Familiarity with LLM agents, AI-powered automation, model governance, and responsible AI practices.

- Experience developing and launching AI-driven products within enterprise environments.

- Hands-on experience with query languages (SQL preferred) and tools such as Tableau, Snowflake, Azure, Pendo, and FullStory.

- Experience in Home Services trades (HVAC, Plumbing, Electrical) is a plus.

****Ways to Stand Out:

****

- Ability to simplify complex AI concepts into actionable insights for cross-functional teams.

- Hands-on experience with LLM fine-tuning, AI retrieval systems, and advanced AI modeling.

```

### Cover letter

```

Dear Hiring Manager,

I am excited to apply to the Principal Product Manager (Titan Intelligence) role at ServiceTitan. After carefully reviewing the responsibilities of my professional background and goals, I made the easy decision to apply. I bring over 15 years of experience in product management and AI-driven innovation with a track record of launching successful products, integrating automation into enterprise solutions, and driving execution through cross-functional collaboration.

I have a strong background in driving AI-powered product initiatives focusing on delivering tangible user benefits. I believe in the power of teamwork across different functions to achieve great results. In my resume and the following sections, I'll highlight how my skills and experiences can support and enhance ServiceTitan's Titan Intelligence initiatives.

-   **AI-Driven Product Leadership:** I have successfully developed and launched AI and machine learning features by translating complex technical capabilities into intuitive user experiences. I ensure every release delivers tangible value and adoption by aligning engineering innovation with market needs.
-   **Data-Informed Strategy:** I focus on using data, testing new ideas, and listening to user feedback to improve how features perform, all while crafting roadmaps that enhance both customer satisfaction and business growth. I believe in working in quick, iterative cycles so that we can learn a lot and make a meaningful impact along the way.
-   **Scalable AI Solutions:** I've designed and integrated scalable architectures that bring together multiple data sources, ensuring that the AI insights are both effective and dependable. This experience has prepared me to assist Titan Intelligence in expanding its capabilities while keeping the user experience smooth and enjoyable.

My background in AI innovation enables me to develop automation and intelligence solutions that reduce manual tasks and enhance productivity for home service professionals. I take a practical approach to AI adoption, ensuring new features are accessible, intuitive, and immediately beneficial. At ServiceTitan, I aim to unify teams around a shared vision, enhance collaboration, and streamline delivery cycles through iterative development.

Thank you for your time and consideration. I look forward to the conversation.
```

## Sample 4

### Sample Job post

```

## About the job

**What We Do:

**At Autura, we're revolutionizing the towing and recovery industry with cutting-edge software that makes a real difference for our customers. From towing service providers to local governments, our tools—like Towing Management Systems, ARIES Dispatch, and Impound solutions—are designed to simplify operations, boost safety, and drive success. Joining Autura means being part of a team that's passionate about innovation, collaboration, and making the world a little safer and smarter every day. Ready to help us pave the way?

**About the Role:

**The Director of Product Management will be responsible for the development and success of the Private Towing segment of our portfolio, primarily focused on TMSs (Towing Management System). This individual will own the roadmap, prioritize features, and guide the product lifecycle from ideation to launch and beyond.

This is a player-coach role and will mentor and grow two other Product Managers focused on private towing operators. The ideal candidate will have a deep understanding of product management, experience in developing software solutions, and a passion for delivering value to customers through intuitive and impactful products, while positively influencing direct reports and the broader team.

**What You'll Be Doing:

**

- Shape and ensure the delivery of the roadmap for multiple Towing Management Systems

- Collaborate with cross-functional teams, including engineering, design, sales, and marketing, to drive the development and launch of product features

- Gather and analyze customer feedback and market trends to inform product decisions and identify new opportunities for growth and improvement

- Define and understand product requirements, ensuring investments create measurable value

- Manage the product lifecycle, including product release plans, timelines, and post-launch analysis to ensure continuous improvement

- Monitor product performance using key metrics and adjust strategies as needed to meet objectives

- Communicate product plans, progress, and results to stakeholders and senior management

- Engage with customers and users to understand their needs and incorporate their feedback into the product development process

- Champion user-centric design and advocate for the best user experience throughout the product development cycle

**About You:

**

- 5+ years of experience as a Product Manager, preferably in the towing, transportation, or field services industries

- Proven ability to manage and deliver software products from concept to launch

- Technical acumen with an understanding of software development processes

- Excellent analytical skills and the ability to make data-informed decisions

- Outstanding communication and collaboration skills to work effectively with cross-functional teams

- Experience equipping software development teams with the information they need to be successful

- Ability to prioritize tasks and manage multiple projects in a fast-paced environment

- Familiarity with agile development methodologies and tools

- Bias for action and appreciation for in-person customer visits and discovery

- Passion for user experience and building customer-centric products

```

### Cover letter

```

Dear Hiring Manager,

I'm thrilled to be applying for the Director, Product Manager position at Autura. After looking closely at the responsibilities and contrasting them with my background, it felt like a natural fit. With over 15 years in product management, solution engineering, and technology consulting under my belt, I've consistently driven innovation and successfully scaled enterprise solutions. I've led impactful product strategies that create real value and contribute to sustainable business growth, aiming to develop products that not only meet user needs but also propel the company forward.

My expertise lies in strategy, execution, and technical leadership, enabling me to translate market insights into products that solve real-world challenges. Below are key areas where my experience aligns with Autura’s goals:

-   **Strategic Product Roadmapping & Lifecycle Execution:** I have built and executed roadmaps that transform complex challenges into actionable plans, ensuring seamless alignment with customer needs and market trends. My work has consistently led to product launches that exceed adoption and revenue targets.
-   **Cross-Functional Leadership & Agile Execution:** I thrive in fast-paced environments where engineering, design, marketing, and sales must work in lockstep. By fostering alignment and accountability, I drive faster product delivery and ensure teams execute with precision.
-   **Customer-Centric Innovation:** I integrate user insights, data-driven experimentation, and advanced technology solutions to build intuitive, high-value experiences. My approach prioritizes usability and impact, increasing engagement and long-term customer satisfaction.

Autura’s vision to modernize the towing industry with cutting-edge software resonates with my background in driving operational transformation through technology. My ability to blend strategic vision with execution makes me confident that I can help propel Autura’s product success.

Thank you for your time and consideration. I look forward to the conversation.
```

---

# Job Post

```text

{job_post}

```
# Analysis

```json

{analysis}

```
# Company
 

```json

{company}

```
 

# Matching the job post requirements with the candidate's key skills and core competencies

The following are the key skills required in the job post matched with the candidate's skill or core competency that is represented among the accomplishments

```json

{key_skills}

```

# Format instructions

{format_instructions}

Reply only with the JSON document that follows the provided schema. Do not add preamble or commentary.
# Quality checks

These are quality checks that you MUST satisfy:

- All the written information is factual and based on the information provided by the talent.
- The cover letter is based on the key matched skills of core competencies of the candidate
- The cover letter's opening paragraph is 3 to 5 sentences long.
- The cover letter core competencies section has 3 to 5 key strengths or experiences.
- The strengths or experiences highlighted in the cover letter are not repeated in the resume.
- The cover letter value alignment paragraph is 4 to 6 sentences long.
- The cover letter call to action paragraph is 2 to 3 sentences long.
- The cover letter is written following the writing style instructions provided
- The cover letter uses the same patterns found in the samples
- The cover letter uses the same tone, style, and sentence structure as the provided samples.
""")


output_parser = JsonOutputParser(pydantic_object=CoverLetter)


def format_accomplishments(accomplishments,job):
    return [ {
            "accomplishment": f'{o["title"]}: {o["body"]}',
            "skills": o["skills"],
            "core_competencies": o["core_competencies"]
            } for o in accomplishments["accomplishments"][job] ]

def get_cover_letter_prompt():

    # Read input/01_job_post.txt as job_post
    with open('inputs/01_job_post.txt', 'r') as file:
        job_post = file.read()

    # read 02_analysis_result.json as Analysis
    with open('inputs/02_analysis_result.json', 'r') as file:
        data = json.load(file)
        analysis_result = Analysis(**data)

    # read 03_accomplishments.json as accomplishments
    with open('inputs/03_accomplishments.json', 'r') as file:
        accomplishments = json.load(file)

    # read 05_key_skills.json as key_skills
    with open('inputs/05_key_skills.json', 'r') as file:
        data = json.load(file)
        key_skills = TopSkills(**data)

    # read 06_company_data.json as company
    with open('inputs/07_company_data.json', 'r') as file:
        data = json.load(file)
        company = CompanyResearch(**data)

    eightBase_accomplishments = format_accomplishments(accomplishments,"8base")
    appify_accomplishments = format_accomplishments(accomplishments,"Appify")
    everteam_accomplishments = format_accomplishments(accomplishments,"Everteam")
    intalio_accomplishments = format_accomplishments(accomplishments,"Intalio")

    prompt = template.format(
        job_post=job_post,
        analysis=json.dumps(analysis_result.model_dump(),indent=4),
        company=json.dumps(company.model_dump(),indent=4),
        key_skills=json.dumps(key_skills.model_dump(),indent=4),
        eightBase_accomplishments=json.dumps(eightBase_accomplishments,indent=4),
        appify_accomplishments=json.dumps(appify_accomplishments,indent=4),
        everteam_accomplishments=json.dumps(everteam_accomplishments,indent=4),
        intalio_accomplishments=json.dumps(intalio_accomplishments,indent=4),
        format_instructions=output_parser.get_format_instructions())

    # write the prompt to a file called prompts/prompt_01.txt
    with open('prompts/cover_letter.md', 'w') as file:
        file.write(prompt)


def sanitize_filename(filename):
    """
    Sanitizes a string to be used as a valid filename.
    Removes invalid characters and replaces whitespaces with underscores.

    :param filename: The original string to sanitize.
    :return: A sanitized string safe for use as a filename.
    """
    # Remove invalid characters (anything except alphanumerics, dots, underscores, and hyphens)
    sanitized = re.sub(r'[^\w.\-]', '', filename)
    # Replace any sequence of whitespace with an underscore
    sanitized = re.sub(r'\s+', '_', sanitized)
    return sanitized


def render_cover():

    clipboard_content = subprocess.check_output(['pbpaste']).decode('utf-8')
    cover_data = json.loads(clipboard_content)
    with open('inputs/08_cover_letter_data.json', 'w') as file:
        json.dump(cover_data, file, indent=4)    

    # read 02_analysis_result.json as Analysis
    with open('inputs/02_analysis_result.json', 'r') as file:
        data = json.load(file)
        analysis_data = Analysis(**data)



    folder_name = sanitize_filename(f"{analysis_data.company_name}-{analysis_data.job_title}")
    store_path = "/Users/estebanf/Library/CloudStorage/GoogleDrive-esteban.felipe@gmail.com/My Drive/JH 2025"

    # create document for cover letter
    cover = Document("cover.docx")

    for paragraph in cover.paragraphs:
        if '{{content}}' in paragraph.text:
            style = paragraph.style
            paragraph.text = cover_data['opening_paragraph'] + "\n"
            current_paragraph = paragraph
            core_competencies_paragraph = cover.add_paragraph()
            core_competencies_paragraph.style = style
            core_competencies_paragraph.text = cover_data['core_competencies_paragraph']['intro_paragraph'] + '\n'
            current_paragraph._p.addnext(core_competencies_paragraph._p)
            current_paragraph = core_competencies_paragraph
            items_paragraph = cover.add_paragraph()
            items_paragraph.style = style
            current_paragraph._p.addnext(items_paragraph._p)
            current_paragraph = items_paragraph

            for item in cover_data['core_competencies_paragraph']['core_competencies']:        # Add bullet point
                p = current_paragraph._element
                pPr = p.get_or_add_pPr()
                numPr = OxmlElement('w:numPr')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), "0")
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), "2")
                numPr.append(ilvl)
                numPr.append(numId)
                pPr.append(numPr)
                
                # Add label in bold
                run = current_paragraph.add_run(f"{item['title']}: ")
                run.bold = True
                
                # Add details
                current_paragraph.add_run(item['details'])
                
                # Insert new paragraph after the current one
                new_paragraph = cover.add_paragraph()
                new_paragraph.style = style
                current_paragraph._p.addnext(new_paragraph._p)
                current_paragraph = new_paragraph
                
            current_paragraph.text = "\n" + cover_data['value_alignment_paragraph'] + "\n"
            call_to_action_paragraph = cover.add_paragraph()
            call_to_action_paragraph.style = style
            call_to_action_paragraph.text = cover_data['call_to_action_paragraph']
            current_paragraph._p.addnext(call_to_action_paragraph._p)

    savepath_cover = sanitize_filename(f"Esteban_Felipe-{analysis_data.company_name}-COVER.docx")
    # Save the file in the created folder
    cover.save(os.path.join(store_path, folder_name, savepath_cover))

